"""
Generate a filled ALU Research Ethics Application Checklist (Jan 2026) for the
Geospatial Flood-Risk Modelling capstone. Secondary-data / ML study with NO
human participants, so participant-facing items are Not Applicable.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

CHECK = "☒"   # ballot box with X
EMPTY = "☐"   # empty ballot box
NAVY = RGBColor(0x1F, 0x35, 0x64)
RED = RGBColor(0xC0, 0x00, 0x00)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
for s in ("Heading 1", "Heading 2"):
    doc.styles[s].font.color.rgb = NAVY


def H(text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def para(text="", bold=False, italic=False, size=10.5, color=None, after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    return p


def kv_table(rows):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in rows:
        cells = t.add_row().cells
        cells[0].width = Inches(2.7)
        cells[1].width = Inches(3.8)
        rk = cells[0].paragraphs[0].add_run(k)
        rk.bold = True
        cells[1].paragraphs[0].add_run(v)
    return t


def choice(label, answer):
    """One risk-assessment row: marks YES or NO."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label + "  ")
    r.font.size = Pt(10)
    yes = f"{CHECK if answer=='YES' else EMPTY} YES   "
    no = f"{CHECK if answer=='NO' else EMPTY} NO"
    ry = p.add_run(yes)
    ry.bold = answer == "YES"
    rn = p.add_run(no)
    rn.bold = answer == "NO"
    return p


def dcheck(label, answer):
    """PART D row with YES / NO / NOT APPLICABLE."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label + "  ")
    r.font.size = Pt(10)
    opts = [("YES", "YES"), ("NO", "NO"), ("NOT APPLICABLE", "N/A")]
    for full, key in opts:
        mark = CHECK if answer == key else EMPTY
        run = p.add_run(f"{mark} {full}    ")
        run.bold = answer == key
    return p


# ---------------------------------------------------------------- Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = t.add_run("January 2026")
rt.bold = True
rt.font.size = Pt(16)
rt.font.color.rgb = NAVY
t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt2 = t2.add_run("Research Ethics Application Checklist")
rt2.bold = True
rt2.font.size = Pt(20)
rt2.font.color.rgb = NAVY

para("All sections of the form are completed below. As this is a secondary-data, "
     "computational study with no human participants, participant-facing items "
     "(consent, recruitment, withdrawal) are marked Not Applicable with justification.",
     italic=True, after=8)

# ---------------------------------------------------------------- PART A
H("PART A: Summary", level=1)
kv_table([
    ("A1. Campus/Hub/Organisation:", "ALU-Rwanda Campus"),
    ("A2. Faculty/Department", "BSE"),
    ("A3. Project Title:", "Geospatial Flood Risk Modelling using Climate Data and "
                           "Hidden Markov Models in Rwanda"),
    ("A4. Student/Principal Researcher's Name", "Kellen Murerwa"),
    ("A5. Email address (institutional email ONLY):", "k.murerwa1@alustudent.com"),
    ("A6. Supervisor's Name:", "Emmanuel Adjei"),
])

para("Links to Documents:", bold=True, after=2)
para("1.  Research Proposal — Flood_Risk_HMM_Capstone_Proposal.pdf (link inserted in submission form).")
para("2.  Data Collection Instrument — Not applicable in the conventional (survey/interview) "
     "sense: no human participants or primary data collection. Data are acquired automatically "
     "from open secondary datasets via documented API queries. The equivalent 'instrument' is the "
     "data-acquisition specification + variable dictionary (study area, 250 m grid, the ERA5 / "
     "SRTM / OpenStreetMap / Rwanda GeoPortal query parameters and the resulting feature schema), "
     "documented in the proposal Methods chapter and in build_real_dataset.py.")
para("3.  Consent Form — Not applicable: the study involves no human participants, so no "
     "informed-consent process is required.")

# A7
H("A7. Project Summary", level=2)
para("Rationale & aims.", bold=True, after=2)
para("Kigali's Nyabugogo–Nyabarongo corridor floods repeatedly during the rainy seasons, "
     "yet operational risk information is largely static (fixed hazard polygons) and does not "
     "reflect day-to-day rainfall dynamics. This project develops a reproducible machine-learning "
     "framework that estimates a daily Low / Moderate / High flood-pressure state for each cell of "
     "the corridor by fusing rainfall, terrain, hydrology and urban-exposure data, adding a Hidden "
     "Markov Model (HMM) temporal layer to capture day-to-day persistence, and validating the "
     "output against the official flood-risk polygons. The aim is a decision-support / risk-"
     "screening tool that complements—not replaces—official early-warning systems.")
para("Methods, procedure & tasks.", bold=True, after=2)
bullet("Build a 250 m analysis grid over the study corridor (729 cells).")
bullet("Acquire real, free, keyless data per cell-day for 2018–2024: rainfall totals "
       "(1/3/7/14-day) from the Open-Meteo ERA5 archive; elevation and slope from SRTM; "
       "distance-to-river, road and building density from OpenStreetMap (Overpass); and the "
       "official flood-risk polygons from the Rwanda GeoPortal (geodata.rw, Ministry of Environment).")
bullet("Derive flood-pressure labels (rainfall trigger × terrain susceptibility with an "
       "unobserved drainage latent and noise so the task is not trivially circular).")
bullet("Train and benchmark rainfall-threshold and static-polygon baselines, Logistic Regression, "
       "Decision Tree, Random Forest, XGBoost (with SHAP explanations) and an HMM temporal layer, "
       "evaluated on a temporal split (train 2018–22, validation 2023, test 2024).")
bullet("Validate spatially against the official polygons (enrichment, inside/outside odds ratio) "
       "and ship a Streamlit inspection dashboard.")
para("Participant sample & recruitment.", bold=True, after=2)
para("None. The study uses only open secondary geospatial/climate datasets; there is no human "
     "participant recruitment, no surveys, interviews or user testing.")
para("Location.", bold=True, after=2)
para("Desk-based / computational research conducted at ALU-Rwanda Campus. The geographic study "
     "area is the Nyabugogo–Nyabarongo corridor, Kigali, Rwanda.")
para("For Software Engineering projects:", bold=True, after=2)
bullet("(a) Nature of software: a machine-learning system (supervised classifiers + an HMM "
       "temporal model) with a Streamlit web dashboard for inspection and visualisation.")
bullet("(b) Data sources / reused code: Open-Meteo ERA5 & SRTM APIs, OpenStreetMap/Overpass, "
       "Rwanda GeoPortal flood polygons; open-source libraries (scikit-learn, XGBoost, hmmlearn, "
       "SHAP, GeoPandas, Streamlit). No proprietary or personal data are used.")
bullet("(c) User interaction: the end user is the researcher/analyst (and, conceptually, planners) "
       "who views maps, per-cell risk states and model metrics through the dashboard. The system "
       "ingests only public geospatial/climate data; it collects no personal data from users.")

# ---------------------------------------------------------------- PART B
doc.add_page_break()
H("PART B: Assessment of Risk", level=1)
para("Does your proposed research study involve any of the following?", italic=True, after=6)

risks = [
    ("B1.  Research involves vulnerable groups (children ≤16; those lacking capacity; "
     "dependent/unequal relationships)", "NO"),
    ("B2.  Research involves sensitive topics (sexual, legal/political behaviour, experience of "
     "violence, gender or ethnic status)", "NO"),
    ("B3.  Research involves a significant element of deception", "NO"),
    ("B4.  Research involves access to records of personal or confidential information "
     "(incl. genetic/biological)", "NO"),
    ("B5.  Research involves access to potentially sensitive data through third parties "
     "(organisational/employee data)", "NO"),
    ("B5.1  Research involves collecting/storing/processing Personally Identifiable Information "
     "(PII) or sensitive user data", "NO"),
    ("B5.2  Research project utilises third-party services (cloud hosting, databases, analytics "
     "platforms, APIs)", "YES"),
    ("B6.  Research that could induce psychological stress, anxiety, humiliation or more than "
     "minimal pain", "NO"),
    ("B7.  Research that may place the researcher at risk of psychological or physical harm", "NO"),
    ("B8.  Research conducted off-campus (archive, community centre, filming, interviewing, etc.)", "NO"),
    ("B9.  Research involving invasive interventions (drugs, vigorous exercise, hypnotherapy)", "NO"),
    ("B10.  Research that may adversely impact employment or social standing", "NO"),
    ("B11.  Research that may lead to 'labelling' by researcher or participant", "NO"),
    ("B12.  Research involving collection of human tissue, blood or biological samples", "NO"),
    ("B13.  Research using procedures that may interact with a pre-existing medical condition", "NO"),
    ("B14.  Research requiring potentially hazardous equipment or environments", "NO"),
    ("B15.  Research requiring ethical approval from another source (e.g. prisons)", "NO"),
    ("B16.  Research requiring permissions from another source (schools, businesses, copyright "
     "holder, commercial sources)", "NO"),
    ("B17.  Research involves development/use of software that collects or processes personal data", "NO"),
    ("B18.  Research involves the use of artificial intelligence / machine-learning models", "YES"),
    ("B19.  Research uses external datasets", "YES"),
    ("B20.  Research involves reusing external code or datasets", "YES"),
    ("B21.  Research involves human-computer interaction (e.g. user testing)", "NO"),
]
for label, ans in risks:
    choice(label, ans)

# Section B explanation
para("")
para("Explanation of 'YES' answers (ethical conduct):", bold=True, after=4)
para("B5.2 — Third-party services / APIs. The project queries only open, public, keyless "
     "data services: Open-Meteo (ERA5 reanalysis & SRTM elevation), the OpenStreetMap Overpass API, "
     "and the Rwanda GeoPortal (geodata.rw). No personal data, accounts or credentials are involved. "
     "API responses are cached locally so the pipeline is reproducible and minimises repeat load on "
     "the services. Source code is stored in a private, university-sanctioned Git repository with "
     "access limited to the researcher and supervisor.")
para("B18 / B19 / B20 — AI/ML models, external datasets and reused code. (i) Licences & "
     "attribution: rainfall and elevation are © Open-Meteo (ERA5/Copernicus; SRTM/NASA); map "
     "data © OpenStreetMap contributors (ODbL); flood polygons © Rwanda GeoPortal / Ministry "
     "of Environment. All are openly licensed for research use and are attributed in the proposal, "
     "code and dashboard; open-source libraries are used under their respective licences. (ii) No "
     "personal data: the datasets are environmental/geospatial; no PII is collected, so no consent "
     "process is required. (iii) Bias, fairness & accountability: models are evaluated on a strict "
     "temporal split (train 2018–22, validation 2023, test 2024) against transparent baselines; feature "
     "attributions (SHAP) and spatial validation against official polygons provide interpretability "
     "and guard against spurious or unfair concentration of predicted risk. Outputs are clearly "
     "framed as a derived 'flood-pressure' risk-screening signal—not a confirmed flood "
     "forecast—and are positioned to complement, not replace, official early-warning systems. "
     "(iv) Data security: all materials are stored on the researcher's ALU Google Drive and a "
     "private Git repository, accessible only to the researcher and supervisor.")

# ---------------------------------------------------------------- PART C
doc.add_page_break()
H("PART C: Data Management and Storage", level=1)
para("During data collection.", bold=True, after=2)
para("This study collects no personal data and no paper forms. All data are open secondary "
     "geospatial/climate datasets acquired via API and cached locally. Source code and derived "
     "datasets are stored in a private, university-sanctioned Git repository and on the researcher's "
     "ALU Google Drive, with access restricted to the researcher and supervisor (the research team). "
     "No data are stored on unencrypted USB sticks or external drives. As there are no participants, "
     "no participant identifiers, anonymisation or recording-destruction steps are required.")
para("Data storage and retention.", bold=True, after=2)
para("All digital materials relating to the study (datasets, model outputs, code and documents) are "
     "stored on the researcher's ALU Google Drive and the private repository for a period of 3 months "
     "following graduation and will then be securely destroyed, in line with ALU policy. The "
     "underlying source datasets remain publicly available from their original open providers.")

# ---------------------------------------------------------------- PART D
H("PART D: Final checklist and declaration", level=1)
dchecks = [
    ("D1.  Have you followed the guidance provided by your subject area?", "YES"),
    ("D2.  Have you included any adverts or letters to recruit participants?", "N/A"),
    ("D3.  Are timescales for participant withdrawal clear, specific and explained?", "N/A"),
    ("D4.  Have you made the mechanisms for withdrawing from the study clear?", "N/A"),
    ("D5.  Have you included your Consent Form?", "N/A"),
    ("D6.  Have you made clear the research is an undergraduate project at ALU?", "N/A"),
    ("D7.  Will you provide participants with your supervisor's name and contact details?", "N/A"),
    ("D8.  Will you provide participants with your name and University email?", "N/A"),
    ("D9.  Have you included copies of all measures/tests/questionnaires and the research "
     "proposal in your application?", "YES"),
    ("D10.  Have you included approval/permissions from other bodies (if applicable)?", "N/A"),
    ("D11.  Have you made clear that personal/sensitive data will be anonymised and stored "
     "securely at ALU for 1 year?", "N/A"),
    ("D12.  Will you make clear who has access to the raw and summarised data?", "YES"),
    ("D13.  Will you make clear that anonymised data may form part of a publication?", "N/A"),
    ("D14.  Have you checked all participant materials for spelling/grammatical errors?", "N/A"),
    ("D15.  Confirm participants will not be subject to undue/disproportionate incentives.", "N/A"),
    ("D16.  Have you identified all third-party software/libraries/APIs/OSS, verified licences "
     "and outlined compliance?", "YES"),
    ("D17.  (Software) Have you included a plan for secure handling of user data during testing?", "YES"),
    ("D18.  (Software) If using external code, have you checked licences and provided attribution?", "YES"),
    ("D19.  (Software) If using ML models / external datasets, have you considered bias, fairness "
     "and accountability and documented them?", "YES"),
]
for label, ans in dchecks:
    dcheck(label, ans)

para("")
para("Notes on Part D answers.", bold=True, after=2)
para("Items D2–D8, D10, D11, D13–D15 are Not Applicable because the study has no human "
     "participants and collects no personal data. D9: the research proposal is included; there are "
     "no questionnaires or tests. D12: raw and summarised data are accessible only to the researcher "
     "and supervisor and derive from public open datasets. D17: the dashboard processes only public "
     "geospatial data and handles no personal user data. D16/D18/D19: licences, attribution and "
     "bias/fairness considerations are documented in the proposal and repository (see Part B).")

# ---------------------------------------------------------------- Supervisor Authorisation
doc.add_page_break()
H("Supervisor Authorisation", level=1)
para("Your application for research ethics review and clearance will NOT be reviewed by the "
     "research ethics committee if your supervisor has not authorised and signed this form.",
     bold=True, color=RED)
para("TO BE COMPLETED BY THE STUDY SUPERVISOR", bold=True, after=4)
para("I confirm that the student/lead investigator has had appropriate health and safety training.")
para(f"{EMPTY} YES     {EMPTY} NO     {EMPTY} N/A")
para("I confirm that I have checked the form and the related materials and agree that, in addition "
     "to general ethical issues, the matters listed above have been specifically addressed.")
para(f"{EMPTY} YES     {EMPTY} NO")
para("Do not submit this form without your supervisor's approval and signature!", bold=True, color=RED)

para("")
sig = doc.add_table(rows=0, cols=2)
sig.style = "Table Grid"
for left, right in [
    ("Signature of Study Supervisor(s):", "______________________________"),
    ("Name(s):", "Emmanuel Adjei"),
    ("Signature of Student / lead investigator:", "______________________________"),
    ("Name(s):", "Kellen Murerwa"),
    ("Date:", "______________________________"),
]:
    cells = sig.add_row().cells
    cells[0].paragraphs[0].add_run(left).bold = True
    cells[1].paragraphs[0].add_run(right)

out = "Research_Ethics_Application_FILLED.docx"
doc.save(out)
print("saved", out)
