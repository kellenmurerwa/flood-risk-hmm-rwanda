"""Generate the full Flood-Risk HMM capstone proposal as a .docx file."""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
FIG = os.path.join(HERE, "figures")
OUT = os.path.join(os.path.dirname(HERE), "Flood_Risk_HMM_Capstone_Proposal.docx")

NAVY = RGBColor(0x0B, 0x3D, 0x91)
LINKBLUE = RGBColor(0x0B, 0x3D, 0x91)

TITLE = ("Geospatial Flood Risk Modelling Using Climate Data and Hidden "
         "Markov Models in Rwanda")

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(6)


def set_run(run, *, bold=False, italic=False, size=None, color=None, font=None):
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if font is not None:
        run.font.name = font
    return run


def add_para(text="", *, bold=False, italic=False, align=None, size=12,
             color=None, space_before=0, space_after=6, indent_first=0,
             line_spacing=1.5):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent_first:
        pf.first_line_indent = Cm(indent_first)
    if text:
        run = p.add_run(text)
        set_run(run, bold=bold, italic=italic, size=size, color=color,
                font="Times New Roman")
    return p


def h1(text, page_break_before=True):
    if page_break_before:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.15
    r = p.add_run(text.upper())
    set_run(r, bold=True, size=14, color=NAVY, font="Times New Roman")
    return p


def h2(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    r = p.add_run(text)
    set_run(r, bold=True, size=13, color=NAVY, font="Times New Roman")
    return p


def h3(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15
    r = p.add_run(text)
    set_run(r, bold=True, italic=True, size=12, font="Times New Roman")
    return p


def justified(text, indent=0.6):
    add_para(text, align="justify", indent_first=indent)


def bullet(text, level=0):
    style_name = "List Bullet" if level == 0 else "List Bullet 2"
    try:
        p = doc.add_paragraph(style=style_name)
    except KeyError:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1 + level)
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(3)
    r = p.add_run(text)
    set_run(r, font="Times New Roman", size=12)
    return p


def numbered(text):
    try:
        p = doc.add_paragraph(style="List Number")
    except KeyError:
        p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run(r, font="Times New Roman", size=12)
    return p


def add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "23")  # 11.5pt
    rPr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0B3D91")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_figure(path, caption, width_in=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run(r, italic=True, size=11, font="Times New Roman")
    cap.paragraph_format.space_after = Pt(12)


def add_table(headers, rows, widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        set_run(r, bold=True, size=11, color=RGBColor(0xFF, 0xFF, 0xFF),
                font="Times New Roman")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "0B3D91")
        tcPr.append(shd)
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            set_run(run, size=10.5, font="Times New Roman")
            p.paragraph_format.space_after = Pt(2)
    if widths_cm:
        for row in table.rows:
            for i, w in enumerate(widths_cm):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def add_table_caption(text):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(text)
    set_run(r, italic=True, size=11, font="Times New Roman")
    cap.paragraph_format.space_after = Pt(6)


# ---------------------------------------------------------------------------
# COVER PAGE
# ---------------------------------------------------------------------------
add_para("African Leadership University", align="center", bold=True, size=20,
         color=NAVY, space_before=72, space_after=6)
add_para("School of Science and Technology", align="center", italic=True,
         size=13, space_after=72)

add_para(TITLE, align="center", bold=True, size=18, color=NAVY, space_after=12,
         line_spacing=1.25)

add_para("Capstone Project Proposal", align="center", italic=True, size=14,
         space_before=48, space_after=120)

add_para("BSc. in Software Engineering", align="center", bold=True, size=14,
         space_after=6)
add_para("Kellen Murerwa", align="center", size=13, space_after=24)
add_para("Supervisor: Emmanuel Adjei", align="center", bold=True, size=13,
         space_after=120)
add_para("24th May 2026", align="center", bold=True, size=12)

# ---------------------------------------------------------------------------
# TABLE OF CONTENTS
# ---------------------------------------------------------------------------
doc.add_page_break()
add_para("TABLE OF CONTENTS", align="center", bold=True, size=14, color=NAVY,
         space_after=12)

toc = [
    ("ABSTRACT", "i"),
    ("LIST OF TABLES", "ii"),
    ("LIST OF FIGURES", "iii"),
    ("LIST OF ACRONYMS AND ABBREVIATIONS", "iv"),
    ("", ""),
    ("CHAPTER ONE: INTRODUCTION", "1"),
    ("    1.1 Introduction and Background", "1"),
    ("    1.2 Problem Statement", "3"),
    ("    1.3 Project Main Objective", "4"),
    ("        1.3.1 Specific Objectives (SMART)", "4"),
    ("    1.4 Research Questions", "5"),
    ("    1.5 Project Scope", "6"),
    ("    1.6 Significance and Justification", "7"),
    ("    1.7 Research Budget", "8"),
    ("    1.8 Research Timeline", "9"),
    ("", ""),
    ("CHAPTER TWO: LITERATURE REVIEW", "10"),
    ("    2.1 Introduction", "10"),
    ("    2.2 Historical Background of the Research Topic", "10"),
    ("    2.3 Overview of Existing Systems", "11"),
    ("    2.4 Review of Related Work", "13"),
    ("    2.5 Summary of the Reviewed Literature", "16"),
    ("    2.6 Strengths and Weaknesses of the Existing Systems", "17"),
    ("        2.6.1 Strengths of Current Approaches", "17"),
    ("        2.6.2 Weaknesses and Limitations", "18"),
    ("        2.6.3 Comparison of this System with Current Approaches", "19"),
    ("    2.7 General Comment and Conclusion", "20"),
    ("", ""),
    ("CHAPTER THREE: SYSTEM ANALYSIS AND DESIGN", "21"),
    ("    3.1 Introduction", "21"),
    ("    3.2 Research Design", "21"),
    ("        3.2.1 Model of Development Framework", "22"),
    ("    3.3 Requirements Specification", "23"),
    ("        3.3.1 Functional Requirements", "23"),
    ("        3.3.2 Non-Functional Requirements", "24"),
    ("    3.4 System Architecture", "25"),
    ("        3.4.1 Data Definition and Acquisition", "26"),
    ("        3.4.2 Machine Learning and HMM Pipeline", "27"),
    ("        3.4.3 Model Comparison and Architecture", "28"),
    ("    3.5 UML and Modelling Diagrams", "29"),
    ("        3.5.1 Use Case Diagram", "29"),
    ("        3.5.2 Class Diagram", "30"),
    ("        3.5.3 Entity-Relationship Diagram", "31"),
    ("        3.5.4 Sequence Diagram", "32"),
    ("        3.5.5 HMM State-Transition Model", "33"),
    ("    3.6 Development Tools", "34"),
    ("", ""),
    ("REFERENCES", "35"),
]
for label, page in toc:
    if not label:
        add_para("", space_after=4)
        continue
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT, 2)
    r = p.add_run(label)
    set_run(r, size=11.5, font="Times New Roman",
            bold=label.startswith(("CHAPTER", "ABSTRACT", "LIST", "REFERENCES")))
    p.add_run("\t" + page).font.name = "Times New Roman"

# ---------------------------------------------------------------------------
# LIST OF TABLES
# ---------------------------------------------------------------------------
doc.add_page_break()
add_para("LIST OF TABLES", align="center", bold=True, size=14, color=NAVY,
         space_after=12)
tables_list = [
    ("Table 1.1", "Research Budget (USD)", "8"),
    ("Table 1.2", "Project Schedule Summary", "9"),
    ("Table 2.1", "Summary of Reviewed Literature", "16"),
    ("Table 2.2", "Strengths and Weaknesses of Existing Approaches", "18"),
    ("Table 2.3", "Comparative Analysis: Proposed System vs Existing Approaches", "19"),
    ("Table 3.1", "Functional Requirements", "23"),
    ("Table 3.2", "Non-Functional Requirements", "24"),
    ("Table 3.3", "Datasets Within the Scope of the Study", "26"),
    ("Table 3.4", "ML and HMM Pipeline Layers", "27"),
    ("Table 3.5", "Candidate Models and Baselines for Flood-Pressure Modelling", "28"),
    ("Table 3.6", "Development Tools and Justification", "34"),
]
for tnum, ttitle, tpage in tables_list:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT, 2)
    r = p.add_run(f"{tnum}: {ttitle}")
    set_run(r, size=11.5, font="Times New Roman")
    p.add_run("\t" + tpage).font.name = "Times New Roman"

# ---------------------------------------------------------------------------
# LIST OF FIGURES
# ---------------------------------------------------------------------------
doc.add_page_break()
add_para("LIST OF FIGURES", align="center", bold=True, size=14, color=NAVY,
         space_after=12)
figs_list = [
    ("Figure 1", "System Architecture of the Geospatial Flood-Pressure Modelling System", "25"),
    ("Figure 2", "Hidden Markov Model of Daily Flood-Pressure State Transitions", "33"),
    ("Figure 3", "Use Case Diagram", "29"),
    ("Figure 4", "Class Diagram (UML)", "30"),
    ("Figure 5", "Entity-Relationship Diagram (Crow's-Foot Notation)", "31"),
    ("Figure 6", "Sequence Diagram: Daily Flood-Pressure Estimation", "32"),
    ("Figure 7", "Project Gantt Chart (24 May - 28 July 2026)", "9"),
    ("Figure 8", "Iterative Agile-Scrum Development Model", "22"),
]
for fnum, ftitle, fpage in figs_list:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT, 2)
    r = p.add_run(f"{fnum}: {ftitle}")
    set_run(r, size=11.5, font="Times New Roman")
    p.add_run("\t" + fpage).font.name = "Times New Roman"

# ---------------------------------------------------------------------------
# LIST OF ACRONYMS
# ---------------------------------------------------------------------------
doc.add_page_break()
add_para("LIST OF ACRONYMS AND ABBREVIATIONS", align="center", bold=True,
         size=14, color=NAVY, space_after=12)

acronyms = [
    ("AOI", "Area of Interest"),
    ("ALU", "African Leadership University"),
    ("API", "Application Programming Interface"),
    ("AUC", "Area Under the (ROC) Curve"),
    ("CHIRPS", "Climate Hazards Group InfraRed Precipitation with Station data"),
    ("DEM", "Digital Elevation Model"),
    ("ENACTS", "Enhancing National Climate Services"),
    ("ERD", "Entity-Relationship Diagram"),
    ("GIS", "Geographic Information System"),
    ("HMM", "Hidden Markov Model"),
    ("IPCC", "Intergovernmental Panel on Climate Change"),
    ("MINEMA", "Ministry in charge of Emergency Management (Rwanda)"),
    ("ML", "Machine Learning"),
    ("MLOps", "Machine Learning Operations"),
    ("NISR", "National Institute of Statistics of Rwanda"),
    ("OSM", "OpenStreetMap"),
    ("REMA", "Rwanda Environment Management Authority"),
    ("RF", "Random Forest"),
    ("RMSE", "Root Mean Square Error"),
    ("SDG", "Sustainable Development Goal"),
    ("SDLC", "Software Development Life Cycle"),
    ("SRTM", "Shuttle Radar Topography Mission"),
    ("TWI", "Topographic Wetness Index"),
    ("UML", "Unified Modeling Language"),
    ("WGS", "World Geodetic System"),
    ("XGBoost", "Extreme Gradient Boosting"),
]
for ab, full in acronyms:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    r = p.add_run(f"{ab:<10}")
    set_run(r, bold=True, font="Courier New", size=11)
    r2 = p.add_run(full)
    set_run(r2, font="Times New Roman", size=11)

# ---------------------------------------------------------------------------
# ABSTRACT
# ---------------------------------------------------------------------------
doc.add_page_break()
add_para("ABSTRACT", align="center", bold=True, size=14, color=NAVY,
         space_after=12)

abstract = (
    "Rwanda's urban and peri-urban catchments face recurrent flood risk driven by intense "
    "rainfall, steep slopes, informal settlement, poor drainage and land-cover change. "
    "Existing flood-risk resources, official hazard maps and climate platforms, are valuable "
    "but largely static: they identify generally vulnerable zones yet do not show how flood "
    "pressure changes from day to day after rainfall. This research proposes the design and "
    "development of a geospatial machine-learning and Hidden Markov Model (HMM) framework that "
    "fuses real daily rainfall, terrain, hydrology and urban-exposure data to estimate "
    "low, moderate and high flood-pressure states over time for a defined flood-prone corridor "
    "in Kigali, the Nyabugogo-Mpazi catchment. The study follows an iterative agile-scrum "
    "methodology over nine weeks (24 May - 28 July 2026). The selected area is divided into a "
    "100-250 m grid; each cell-day is linked to rainfall accumulation features (1-, 3-, 7- and "
    "14-day totals) from CHIRPS and Meteo Rwanda, terrain features (elevation, slope) from SRTM, "
    "hydrology and exposure features from OpenStreetMap, and intersection with official "
    "flood-risk polygons. A supervised classifier (Random Forest, XGBoost) scores each cell-day "
    "as low, moderate or high flood pressure, and an HMM temporal layer estimates daily state "
    "transitions. The framework is hypothesised to achieve a macro F1-score of at least 0.75, a "
    "high-pressure recall of at least 0.80, and spatial agreement in which at least 70% of "
    "predicted high-pressure cells fall within official flood-risk polygons. The project is "
    "framed honestly as flood-pressure estimation and risk screening rather than confirmed "
    "flood-event prediction, contributing a reproducible, time-aware geospatial workflow for "
    "decision support in data-scarce East African cities."
)
p = add_para(abstract, align="justify")
p.paragraph_format.line_spacing = 1.5
add_para("Keywords: flood risk, Hidden Markov Model, geospatial machine learning, climate "
         "data, CHIRPS, flood-pressure states, Rwanda, Kigali, Nyabugogo, decision support.",
         align="justify", italic=True, space_before=12)

# ---------------------------------------------------------------------------
# CHAPTER 1: INTRODUCTION
# ---------------------------------------------------------------------------
h1("Chapter One: Introduction")

h2("1.1 Introduction and Background")
justified(
    "Rwanda, the 'land of a thousand hills', combines steep topography, dense settlement and "
    "an increasingly variable rainfall regime, a combination that places its towns and cities "
    "at persistent risk of flooding. The Intergovernmental Panel on Climate Change identifies "
    "East Africa as a hotspot for rainfall irregularity and intensifying extreme-precipitation "
    "events (IPCC, 2022), and downscaled projections released by the Rwanda Meteorology Agency "
    "(2022) confirm rising intra-seasonal variability and heavier short-duration storms. When "
    "such storms fall on steep, poorly drained urban catchments, surface runoff concentrates "
    "rapidly, producing inundation, road disruption, property damage and loss of life."
)
justified(
    "The human cost is well documented. In a single month in 2023, heavy rainfall triggered "
    "floods and landslides that killed more than 120 people and damaged hundreds of homes across "
    "northern and western Rwanda, and the Ministry in charge of Emergency Management records "
    "flooding among the most frequent and damaging hazards in its annual disaster statistics "
    "(MINEMA, 2023). In Kigali, flood-prone corridors cluster near streams, wetlands and "
    "low-lying settlements; the Nyabugogo-Mpazi catchment in particular has been studied for "
    "decades as an emblematic urban flood plain where informal settlement and commercial "
    "activity coexist with recurrent inundation (MINEMA, 2023; The New Times, 2026)."
)
justified(
    "A growing body of work shows that geographic information systems (GIS) and machine learning "
    "(ML) can map where flooding is likely. Mind'je et al. (2019) modelled national flood "
    "susceptibility for Rwanda using logistic regression over ten terrain, hydrological and "
    "land-cover factors, while Mind'je et al. (2021) coupled geospatial analysis with "
    "hydrological modelling to simulate peak flow and volume in the Nyabarongo catchment. "
    "Internationally, systematic reviews report that ensemble learners such as Random Forest and "
    "gradient boosting are among the most consistent performers for flood-susceptibility and "
    "flood-risk tasks (Mosavi, Ozturk, & Chau, 2018). These approaches, however, typically "
    "produce a single static map. They answer 'where is generally vulnerable?' but not the "
    "sharper, more operational question this research addresses: given recent rainfall and "
    "the fixed geography of a place, how does its flood pressure move between low, moderate and "
    "high states from one day to the next?"
)

h2("1.2 Problem Statement")
justified(
    "First, existing flood-risk resources for Rwanda are predominantly static. Official hazard "
    "layers and susceptibility maps "
    "(Mind'je et al., 2019) identify vulnerable zones effectively, but they are spatial snapshots "
    "that do not show how flood pressure changes day to day following rainfall events. A "
    "community living beside the "
    "Mpazi drain is 'in a flood-risk zone' every day of the year on a static map, which gives "
    "responders little help in deciding when that zone is most dangerous."
)
justified(
    "Second, climate services provide rainfall information, but rainfall values alone do not "
    "directly explain how a specific urban catchment moves between levels of flood pressure. "
    "Datasets such as CHIRPS (Climate Hazards Center, 2024) and the national ENACTS rainfall products offer "
    "long, gridded daily series, yet translating accumulated rainfall into a per-cell, "
    "time-varying flood-pressure state requires fusing it with terrain, drainage proximity and "
    "exposure, and modelling the temporal dynamics explicitly."
)
justified(
    "Third, although machine-learning flood studies are now common, most classify flood "
    "susceptibility from terrain, land cover and distance-to-river features without a temporal "
    "state model (Khosravi et al., 2018; Ma et al., 2021). There is therefore a documented gap "
    "for a localised, reproducible workflow that integrates real daily climate data, official "
    "flood-risk spatial layers, terrain features and urban-exposure indicators into a single "
    "time-aware model of flood-pressure states."
)
justified(
    "This research proposes a geospatial ML and Hidden Markov Model framework that: (i) divides "
    "a selected flood-prone corridor in Kigali into grid cells linked to daily rainfall and "
    "static spatial features; (ii) derives and classifies low, moderate and high flood-pressure "
    "states with supervised models; and (iii) applies an HMM as a temporal layer that estimates "
    "how those states transition over a daily time step. To remain academically honest, the "
    "system claims to estimate derived flood-pressure states validated against official "
    "flood-risk polygons, not to predict confirmed flood events, unless verified daily incident "
    "records become available."
)

h2("1.3 Project Main Objective")
justified(
    "To develop and evaluate a geospatial machine-learning and Hidden Markov Model framework "
    "that uses real rainfall and spatial data to estimate daily flood-pressure states for a "
    "selected flood-prone area in Rwanda, the Nyabugogo-Mpazi corridor in Kigali, and to "
    "validate those estimates against official flood-risk spatial information."
)

h3("1.3.1 Specific Objectives (SMART)")
numbered(
    "To define a flood-prone study area in Kigali (the Nyabugogo-Mpazi corridor) and prepare a "
    "100-250 m grid-based spatial unit, and to acquire and preprocess daily rainfall (CHIRPS / "
    "Meteo Rwanda, 2018-2024), SRTM elevation, OpenStreetMap geolocation data and official "
    "flood-risk polygons for that area, completed by 21 June 2026."
)
numbered(
    "To engineer daily rainfall accumulation features (1-, 3-, 7- and 14-day totals) and static "
    "geospatial features (elevation, slope, distance-to-river, road and building density, "
    "flood-polygon intersection) into a reproducible grid-cell-by-day modelling table, and to "
    "derive low, moderate and high flood-pressure labels, by 28 June 2026."
)
numbered(
    "To train and benchmark at least three supervised models (a Logistic Regression / Decision "
    "Tree baseline, Random Forest and XGBoost) against a rainfall-threshold baseline and a "
    "static flood-polygon baseline, achieving a macro F1-score of at least 0.75 and a "
    "high-pressure-state recall of at least 0.80 on a held-out period, by 12 July 2026."
)
numbered(
    "To integrate a Hidden Markov Model temporal layer that estimates daily transitions between "
    "flood-pressure states, and to evaluate it on next-step accuracy and spatial agreement, "
    "targeting at least 70% of predicted high-pressure cells falling within official flood-risk "
    "polygons, by 22 July 2026."
)
numbered(
    "To produce flood-pressure risk maps, a reproducible notebook/dashboard for inspecting "
    "results by date and location, and a final report and defence presentation evidencing the "
    "framework's value for decision support, by 28 July 2026."
)

h2("1.4 Research Questions")
numbered(
    "How can daily rainfall data and geospatial features be combined to estimate flood-pressure "
    "states in a selected flood-prone area in Kigali, Rwanda?"
)
numbered(
    "Which geospatial features, such as slope, elevation, river proximity, building density, "
    "road density and official flood-risk intersection, contribute most to flood-pressure "
    "classification?"
)
numbered(
    "How well does a geospatial ML classifier separate low, moderate and high flood-pressure "
    "states compared with a rule-based rainfall-threshold baseline and a static flood-polygon "
    "baseline?"
)
numbered(
    "How does a Hidden Markov Model represent transitions between flood-pressure states over a "
    "daily observation interval, and does temporal smoothing improve next-step stability?"
)
numbered(
    "How closely do predicted high-pressure cells align with official flood-risk polygons or any "
    "available incident records?"
)

h2("1.5 Project Scope")
justified(
    "Geographically, the study is confined to one clearly defined flood-prone catchment, the "
    "Nyabugogo-Mpazi corridor in Kigali, selected for its strong urban flood-risk relevance, "
    "steep slopes, drainage pressure and dense settlement and infrastructure exposure. The area "
    "is small enough to allow grid-based modelling and detailed feature engineering at a 100-250 "
    "m resolution, depending on data resolution and computational capacity."
)
justified(
    "Temporally, the modelling period covers 2018-2024, or another window for which continuous "
    "daily rainfall data is available. The model focuses on derived flood-pressure states rather "
    "than confirmed disaster events. A flood-pressure state is a condition derived from rainfall "
    "accumulation, terrain, proximity to rivers and exposure indicators; a confirmed flood event "
    "requires observed incident data. Where verified, dated and geolocated incident records are "
    "unavailable, claims remain limited to flood-pressure modelling and official flood-risk "
    "polygon agreement."
)
justified(
    "Technically, the deliverable is a reproducible data-and-model pipeline with maps and a "
    "lightweight notebook or dashboard. The project excludes physically based hydrodynamic "
    "engineering models, official early-warning issuance, and real-time sensor networks, which "
    "are noted as future work. The system is positioned as a decision-support and risk-screening "
    "framework that complements, rather than replaces, official warning systems."
)

h2("1.6 Significance and Justification")
justified(
    "For disaster managers and city authorities, the proposed framework moves beyond static "
    "hazard maps toward a time-aware view of flood pressure, helping prioritise inspection, "
    "drainage maintenance and community alerts in the days following heavy rainfall. By grounding "
    "estimates in official flood-risk polygons, it produces outputs that are interpretable "
    "alongside resources already trusted by MINEMA and Kigali City."
)
justified(
    "For the research and software-engineering community, the project contributes a reproducible "
    "integration of supervised geospatial ML with a Hidden Markov temporal layer for an "
    "African urban catchment, an under-explored pairing that most flood-susceptibility studies "
    "omit (Mosavi et al., 2018; Khosravi et al., 2018). Aligned with Sustainable Development "
    "Goals 11 (Sustainable Cities) and 13 (Climate Action) and Rwanda's Vision 2050 resilience "
    "priorities, it offers a low-cost, open-data blueprint transferable to other data-scarce "
    "East African cities."
)
justified(
    "For affected communities, the longer-term significance lies in screening: a transparent, "
    "data-grounded indication of when and where flood pressure is building can support the local, "
    "community-based disaster-risk-reduction approaches that have proven effective in Rwanda."
)

h2("1.7 Research Budget")
justified(
    "The project is deliberately scoped to rely on free-tier, open-data and open-source "
    "components wherever possible. Table 1.1 summarises the costs that must be funded, all of "
    "which are non-substitutable.",
    indent=0
)
add_table(
    ["#", "Item", "Description", "Qty", "Unit Cost (USD)", "Total (USD)"],
    [
        ["1", "Cloud compute / GPU credits", "Google Colab Pro or Kaggle for model training", "2 months", "10.00", "20.00"],
        ["2", "Cloud hosting (Render / Railway free + paid)", "Dashboard + lightweight API during evaluation", "2 months", "7.50", "15.00"],
        ["3", "High-resolution rainfall / data access", "Meteo Rwanda station data beyond free tier", "1", "25.00", "25.00"],
        ["4", "Object storage", "Versioned storage for raster and vector layers", "2 months", "5.00", "10.00"],
        ["5", "Field validation visit to Kigali corridor", "Ground-truth photos, transport, 2 visits", "2", "30.00", "60.00"],
        ["6", "Expert review (GIS / hydrology)", "Consultation on labels and validation", "1", "30.00", "30.00"],
        ["7", "Printing and stationery", "Maps, consent forms, validation scripts", "1", "10.00", "10.00"],
        ["8", "Contingency (10%)", "Unforeseen expenses", "1", "17.00", "17.00"],
    ],
    widths_cm=[1, 4, 5.5, 1.5, 2, 2]
)
add_table_caption("Table 1.1: Research Budget (USD). Estimated total: USD 187.00 (~RWF 234,000 at 1 USD = 1,250 RWF).")

h2("1.8 Research Timeline")
justified(
    "The project executes in nine weeks between 24 May and 28 July 2026. Figure 7 presents the "
    "Gantt chart and Table 1.2 summarises the major milestones.",
    indent=0
)
add_figure(os.path.join(FIG, "fig07_gantt.png"),
           "Figure 7: Project Gantt Chart (24 May - 28 July 2026).")
add_table(
    ["Week", "Dates", "Key Deliverable"],
    [
        ["1-2", "24 May - 6 Jun", "Literature review complete; study area defined; datasets acquired"],
        ["3", "7 Jun - 13 Jun", "Grid created; geospatial feature engineering finalised"],
        ["4", "14 Jun - 20 Jun", "Rainfall features and flood-pressure labels constructed"],
        ["5", "21 Jun - 27 Jun", "Baselines and ML classifiers trained"],
        ["6", "28 Jun - 4 Jul", "Best classifier selected; HMM temporal layer prototyped"],
        ["7", "5 Jul - 11 Jul", "HMM integration; end-to-end pipeline run"],
        ["8", "12 Jul - 18 Jul", "Evaluation, spatial validation and risk mapping"],
        ["9", "19 Jul - 28 Jul", "Results analysis, dashboard, final report, defence prep"],
    ],
    widths_cm=[2, 4, 10]
)
add_table_caption("Table 1.2: Project Schedule Summary.")

# ---------------------------------------------------------------------------
# CHAPTER 2: LITERATURE REVIEW
# ---------------------------------------------------------------------------
h1("Chapter Two: Literature Review")

h2("2.1 Introduction")
justified(
    "This chapter reviews the geospatial, hydrological and machine-learning literature relevant "
    "to flood-risk modelling, with particular attention to Rwandan and East African contexts and "
    "to temporal state modelling using Markov-based methods. The search targeted peer-reviewed "
    "journal articles, conference papers and credible institutional reports published mainly "
    "between 2018 and 2026, drawn from ScienceDirect (Elsevier), MDPI, Springer, Wiley, IEEE "
    "Xplore, Google Scholar and government and agency repositories including MINEMA, NISR and "
    "Meteo Rwanda. The search combined the keyword clusters {flood OR inundation OR flood "
    "susceptibility} AND {machine learning OR random forest OR XGBoost OR hidden Markov model} "
    "AND {Rwanda OR Kigali OR East Africa OR urban catchment}. After screening, the most relevant "
    "studies were grouped into flood-susceptibility mapping, climate and terrain data sources, "
    "and temporal Markov modelling, of which fourteen are discussed in depth below."
)

h2("2.2 Historical Background of the Research Topic")
justified(
    "Flood-risk analysis evolved from physically based hydrological and hydraulic models, which "
    "remain authoritative but data- and expertise-intensive, toward data-driven approaches that "
    "exploit increasingly available satellite and open geospatial data. The availability of "
    "reprocessed global terrain data such as NASADEM (NASA JPL, 2020) and of long, "
    "quasi-global daily rainfall estimates such as CHIRPS (Climate Hazards Center, 2024) made it "
    "feasible to characterise flood drivers consistently even in sparsely gauged regions like "
    "Rwanda."
)
justified(
    "In parallel, machine learning matured into a mainstream tool for flood prediction and "
    "susceptibility mapping. A widely cited review by Mosavi, Ozturk and Chau (2018) documented "
    "the rapid uptake of ML for flood prediction and concluded that ensemble and hybrid methods "
    "deliver strong, cost-effective performance. On the temporal side, Markov and Hidden Markov "
    "Models have a long pedigree in hydrology and climatology: Stoner and Economou (2020) "
    "developed an advanced HMM for rainfall time series in which an unobserved state governs the "
    "precipitation process, while Xie, Jiang and Sainju (2018) generalised the HMM from a "
    "one-dimensional sequence to a two-dimensional map for flood-extent mapping. The intersection "
    "of these trajectories, "
    "geospatial ML for spatial flood pressure and HMMs for its temporal evolution, remains "
    "under-explored for African urban catchments, motivating the present research."
)

h2("2.3 Overview of Existing Systems")
justified(
    "Existing solutions in the problem space fall into four broad categories."
)
bullet("Official flood-risk and hazard maps (e.g., Rwanda GeoPortal layers and susceptibility "
       "maps from Mind'je et al., 2019). These identify generally vulnerable zones and are "
       "authoritative for planning, but they are spatial snapshots rather than daily temporal "
       "models.")
bullet("Climate and rainfall platforms (Meteo Rwanda / ENACTS Maproom; CHIRPS, Climate Hazards "
       "Center, 2024). These provide rainfall and temperature data and seasonal monitoring, but do not "
       "translate daily rainfall accumulation into flood-pressure states for a specific urban "
       "corridor.")
bullet("Hydrological and hydraulic models for Rwanda (e.g., HEC-HMS peak-flow simulation in the "
       "Nyabarongo catchment, Mind'je et al., 2021). These are physically rigorous but require "
       "specialist calibration and are not designed as reproducible, feature-based ML workflows.")
bullet("Machine-learning flood-susceptibility studies (Khosravi et al., 2018; Ma et al., 2021; "
       "Mosavi et al., 2018). These classify susceptibility from terrain and hydrological "
       "features with high accuracy but typically output a single static map with no temporal "
       "state transitions.")

h2("2.4 Review of Related Work")
justified(
    "Mind'je et al. (2019) modelled flood susceptibility and hazard perception in Rwanda using "
    "logistic regression over a national flood inventory of 153 historical flood locations and "
    "ten predictors, including elevation, slope, distance from rivers and roads, NDVI, the "
    "Topographic Wetness Index and rainfall. The study confirmed that terrain and hydrological "
    "proximity dominate flood susceptibility in Rwanda, directly informing the feature design of "
    "the present work, but produced a static susceptibility surface without temporal dynamics."
)
justified(
    "Mind'je et al. (2021) integrated geospatial analysis with HEC-GeoHMS and HEC-HMS "
    "hydrological modelling to simulate peak flow and volume for the Nyabarongo catchment, "
    "demonstrating the value of DEM-derived physiographic parameters. Their physically based "
    "approach is complementary to this proposal: it motivates the terrain and drainage features "
    "used here while highlighting the data and calibration burden that a lighter ML pipeline "
    "seeks to avoid."
)
justified(
    "Most directly relevant, Nzabonantuma et al. (2025) mapped flood susceptibility in the "
    "Nyabarongo Catchment of Rwanda using Random Forest, support vector machine and XGBoost "
    "models with SHapley Additive exPlanations (SHAP), over conditioning factors almost identical "
    "to those proposed here (elevation, slope, distance to river and road, topographic wetness "
    "index and rainfall). Random Forest performed best, with an AUC of 0.968 and an F1-score of "
    "0.92, providing a strong Rwandan precedent for the exact classifier stack and explainability "
    "approach this project adopts, while still producing a static susceptibility surface rather "
    "than a temporal state model. Complementing this, Hahirwabasenga et al. (2024) reviewed "
    "flooding in the Sebeya catchment and emphasised that limited data availability and quality "
    "are the central obstacles to flood-risk management in Rwanda, reinforcing this project's "
    "reliance on open, reproducible data sources and its honest framing of derived flood-pressure "
    "states."
)
justified(
    "Papaioannou et al. (2019) demonstrated flood-inundation mapping at an ungauged basin by "
    "coupling hydrometeorological and hydraulic models, showing that flood extent can be "
    "reconstructed even where gauge data are scarce, a constraint that also characterises Kigali's "
    "urban catchments. Their work underscores the value of fusing rainfall and terrain data for "
    "flood characterisation and motivates the data-driven, feature-based approach adopted here as "
    "a lighter alternative to full hydrodynamic modelling."
)
justified(
    "Mosavi, Ozturk and Chau (2018) reviewed two decades of machine-learning flood-prediction "
    "research and found that ensemble methods, data decomposition and model optimisation are the "
    "most effective strategies, with Random Forest and gradient boosting recurring as strong "
    "performers. This review grounds the choice of Random Forest and XGBoost as the core "
    "classifiers benchmarked in Chapter 3."
)
justified(
    "Khosravi et al. (2018) compared decision-tree algorithms for flash-flood susceptibility at "
    "the Haraz watershed in northern Iran, reporting strong discrimination (high AUC) and "
    "confirming the suitability of tree-based learners for flood tasks with terrain and "
    "hydrological predictors. Ma et al. (2021) developed an XGBoost-based method for flash-flood "
    "risk assessment, achieving a testing accuracy of about 0.84 and outperforming a "
    "least-squares support-vector-machine baseline, evidence that gradient boosting is a sound "
    "primary model for this proposal."
)
justified(
    "Long et al. (2024) systematically reviewed urban flood-resilience evaluation, mapping the "
    "frameworks, methods and limitations of contemporary approaches and noting a recurring "
    "reliance on static indicators with limited treatment of how risk evolves over time. Their "
    "synthesis reinforces the gap this proposal targets, the absence of a temporal state model, "
    "and the centrality of a well-constructed labelling strategy, which here is grounded in "
    "rainfall accumulation and official flood-risk polygons in the absence of dense incident "
    "records."
)
justified(
    "On the temporal side, Stoner and Economou (2020) developed an advanced Hidden Markov Model "
    "for rainfall time series in which an unobserved state, evolving as a Markov chain, governs "
    "the precipitation process. Their formulation is a direct conceptual ancestor of the "
    "flood-pressure HMM proposed here, in which the hidden state is the latent flood-pressure "
    "condition of a grid cell. Xie, Jiang and Sainju (2018) extended the Hidden Markov Model to a "
    "two-dimensional geographical hidden Markov tree for flood-extent mapping, reporting marked "
    "gains over standard classifiers and demonstrating that HMM machinery, including the "
    "forward-backward, Viterbi and Baum-Welch algorithms used to estimate transition matrices and "
    "decode state sequences, transfers effectively to spatial flood problems."
)
justified(
    "Finally, the human and institutional context is documented by MINEMA (2023), whose annual "
    "disaster statistics rank floods among Rwanda's most frequent and damaging hazards, and by "
    "Meteo Rwanda (2022), whose downscaled projections anticipate heavier short-duration rainfall. "
    "Together these confirm both the urgency and the data availability that make a localised, "
    "time-aware flood-pressure model timely and feasible."
)

h2("2.5 Summary of the Reviewed Literature")
add_table(
    ["#", "Author(s) & Year", "Focus", "Key Finding", "Limitation"],
    [
        ["1", "Nzabonantuma et al. (2025)", "RF/SVM/XGBoost+SHAP susceptibility, Nyabarongo, Rwanda", "RF best (AUC 0.968, F1 0.92)", "Static susceptibility; no temporal model"],
        ["2", "Hahirwabasenga et al. (2024)", "Review of flooding, Sebeya catchment, Rwanda", "Data scarcity is the key FRM obstacle", "Review; no predictive model"],
        ["3", "Mind'je et al. (2019)", "Logistic-regression flood susceptibility, Rwanda", "Terrain & hydrology dominate susceptibility", "Static map; no temporal dynamics"],
        ["4", "Mind'je et al. (2021)", "HEC-HMS peak-flow simulation, Nyabarongo", "DEM-driven physiographic modelling works", "Heavy calibration; not ML/reproducible"],
        ["5", "Papaioannou et al. (2019)", "Coupled inundation mapping, ungauged basin", "Flood extent reconstructed without gauges", "Hydraulic; data/compute intensive"],
        ["6", "Mosavi et al. (2018)", "Review of ML flood prediction", "Ensembles and hybrids most effective", "No single deployed local workflow"],
        ["7", "Khosravi et al. (2018)", "Decision trees, flash-flood, Iran", "Tree learners give high AUC", "Different region; static susceptibility"],
        ["8", "Ma et al. (2021)", "XGBoost flash-flood risk", "Accuracy ~0.84; beats SVM baseline", "No temporal state model"],
        ["9", "Long et al. (2024)", "Review of urban flood-resilience evaluation", "Maps frameworks/methods; static indicators dominate", "Little temporal-evolution modelling"],
        ["10", "Stoner & Economou (2020)", "Advanced HMM for rainfall time series", "Hidden state governs precipitation", "Rainfall focus, not flood pressure"],
        ["11", "Xie et al. (2018)", "Geographical hidden Markov tree, flood extent", "HMM beats classifiers spatially (F ~95%)", "Imagery-based; not state-over-time"],
        ["12", "Feng et al. (2025)", "Agent-based flood-resilience modelling", "Couples household & institutional response", "Simulation; not observational ML"],
        ["13", "NASA JPL (2020)", "NASADEM global DEM", "Consistent elevation/slope source", "30 m may smooth fine drainage"],
        ["14", "MINEMA (2023)", "Rwanda disaster statistics", "Floods among top hazards nationally", "Aggregated; limited daily geolocation"],
    ],
    widths_cm=[0.8, 3.3, 3.2, 4.1, 3.6]
)
add_table_caption("Table 2.1: Summary of Reviewed Literature.")

h2("2.6 Strengths and Weaknesses of the Existing Systems")

h3("2.6.1 Strengths of Current Approaches")
bullet("Official hazard maps and susceptibility studies provide authoritative, planning-grade "
       "spatial information and are trusted by national institutions (Mind'je et al., 2019).")
bullet("Open climate and terrain data (CHIRPS, SRTM, OpenStreetMap) make reproducible, low-cost "
       "feature engineering feasible even in data-scarce settings (Climate Hazards Center, 2024; "
       "NASA JPL, 2020).")
bullet("Machine-learning flood models, especially Random Forest and XGBoost, achieve high "
       "discrimination on terrain and hydrological features (Mosavi et al., 2018; Ma et al., "
       "2021).")
bullet("Hidden Markov Models offer a principled, well-understood framework for modelling "
       "transitions between unobserved states over time (Xie et al., 2018; Stoner & Economou, "
       "2020).")

h3("2.6.2 Weaknesses and Limitations")
bullet("Official flood-risk resources are largely static and do not show how flood pressure "
       "changes from day to day after rainfall.")
bullet("Rainfall platforms supply climate data but do not convert daily accumulation into "
       "per-cell flood-pressure states for a specific corridor.")
bullet("Most ML flood studies output a single susceptibility map and omit any temporal state "
       "model of how risk evolves.")
bullet("Hydrological engineering models are rigorous but data- and expertise-intensive, limiting "
       "rapid, reproducible reuse.")
bullet("No reviewed system simultaneously combines: (a) real daily climate data, (b) official "
       "flood-risk spatial layers, (c) terrain and urban-exposure features, and (d) an explicit "
       "temporal state model, in one reproducible workflow for a Rwandan urban catchment.")
add_table(
    ["Approach", "Strength", "Weakness"],
    [
        ["Official hazard / susceptibility maps", "Authoritative, planning-grade", "Static; no daily temporal view"],
        ["Climate / rainfall platforms", "Long, consistent daily data", "No per-cell flood-pressure state"],
        ["Hydrological (HEC-HMS) models", "Physically rigorous peak-flow", "Calibration-heavy; not reproducible ML"],
        ["ML susceptibility (RF / XGBoost)", "High spatial discrimination", "Single static map; no transitions"],
        ["Markov / HMM temporal models", "Principled state-transition modelling", "Rarely fused with geospatial flood features"],
    ],
    widths_cm=[4.5, 5.5, 5.5]
)
add_table_caption("Table 2.2: Strengths and Weaknesses of Existing Approaches.")

h3("2.6.3 Comparison of this System with Current Approaches")
add_table(
    ["Capability", "Hazard map", "Rainfall platform", "ML susceptibility", "Proposed System"],
    [
        ["Uses real daily climate data", "No", "Yes", "Partial", "Yes"],
        ["Per-cell spatial resolution (grid)", "Partial", "No", "Yes", "Yes"],
        ["Daily temporal state modelling (HMM)", "No", "No", "No", "Yes"],
        ["Integrates official flood-risk polygons", "Yes", "No", "Partial", "Yes"],
        ["Urban-exposure features (roads/buildings)", "Partial", "No", "Partial", "Yes"],
        ["Reproducible open-data workflow", "No", "Partial", "Partial", "Yes"],
        ["Honest risk-screening framing", "Yes", "n/a", "Partial", "Yes"],
    ],
    widths_cm=[5, 2.4, 2.8, 2.8, 3]
)
add_table_caption("Table 2.3: Comparative Analysis: Proposed System vs Existing Approaches.")

h2("2.7 General Comment and Conclusion")
justified(
    "The reviewed literature converges on three robust findings: open climate and terrain data "
    "are sufficient to characterise flood drivers in data-scarce regions; ensemble ML models "
    "classify flood susceptibility with high accuracy from terrain and hydrological features; and "
    "Hidden Markov Models provide a principled means of modelling transitions between unobserved "
    "states over time. However, no single system integrates these findings into a deployed, "
    "reproducible, time-aware flood-pressure model for a Rwandan urban catchment, and most flood "
    "ML studies stop at a static susceptibility map. The proposed system targets precisely this "
    "gap, contributing both a software-engineering deliverable, a reproducible geospatial ML and "
    "HMM pipeline, and an empirical evaluation of the pairing for the Nyabugogo-Mpazi corridor in "
    "Kigali, framed honestly as flood-pressure estimation validated against official flood-risk "
    "spatial information."
)

# ---------------------------------------------------------------------------
# CHAPTER 3: SYSTEM ANALYSIS AND DESIGN
# ---------------------------------------------------------------------------
h1("Chapter Three: System Analysis and Design")

h2("3.1 Introduction")
justified(
    "This chapter describes the methodological and architectural choices that translate the "
    "objectives of Chapter 1 and the gaps identified in Chapter 2 into an implementable software "
    "system. The design approach combines (a) an iterative agile-scrum SDLC adapted to the "
    "nine-week project window, (b) a layered, reproducible data-and-model pipeline, and (c) UML "
    "and HMM modelling to specify behaviour, data and temporal dynamics."
)

h2("3.2 Research Design")
justified(
    "The research adopts a design-science methodology, treating the geospatial flood-pressure "
    "modelling pipeline as the central artefact whose construction and evaluation generate the "
    "knowledge contribution (vom Brocke, Hevner, & Maedche, 2020). Both quantitative and qualitative evidence are "
    "collected: quantitative model-performance metrics (macro F1-score, balanced accuracy, "
    "high-pressure recall and precision) and spatial-agreement metrics from validation against "
    "official flood-risk polygons, complemented by qualitative interpretation of feature "
    "importance, transition matrices and produced maps."
)

h3("3.2.1 Model of Development Framework")
justified(
    "An iterative agile-scrum SDLC is adopted in two-week sprints. Each sprint completes a "
    "Plan-Design-Build-Test-Review cycle, with the supervisor acting as product owner. The model "
    "favours rapid feedback incorporation, which is critical given the compressed timeline and "
    "the exploratory, data-driven nature of flood-pressure label construction and model "
    "evaluation."
)
add_figure(os.path.join(FIG, "fig08_sdlc.png"),
           "Figure 8: Iterative Agile-Scrum Development Model Adapted for the Project.")

h2("3.3 Requirements Specification")

h3("3.3.1 Functional Requirements")
add_table(
    ["ID", "Requirement", "Priority"],
    [
        ["FR-01", "The system shall ingest daily rainfall data (CHIRPS / Meteo Rwanda) for the selected area of interest and time period.", "Must"],
        ["FR-02", "The system shall generate a uniform grid (100-250 m) over the AOI with stable grid IDs and centroid coordinates.", "Must"],
        ["FR-03", "The system shall derive terrain features (elevation, slope) per grid cell from SRTM DEM data.", "Must"],
        ["FR-04", "The system shall compute hydrology and exposure features (distance to river/drainage, road density, building density) per cell from OpenStreetMap.", "Must"],
        ["FR-05", "The system shall compute rolling rainfall accumulation features (1-, 3-, 7- and 14-day totals) for each grid-cell-by-day record.", "Must"],
        ["FR-06", "The system shall derive low, moderate and high flood-pressure labels from rainfall accumulation and official flood-risk polygon information.", "Must"],
        ["FR-07", "The system shall train, evaluate and persist a supervised classifier returning a predicted flood-pressure state and probability per cell-day.", "Must"],
        ["FR-08", "The system shall estimate an HMM transition matrix and next-step state probabilities over the daily time step.", "Should"],
        ["FR-09", "The system shall produce flood-pressure maps and compute spatial agreement against official flood-risk polygons.", "Should"],
        ["FR-10", "The system shall expose a notebook or dashboard to inspect predictions by date and grid cell.", "Could"],
    ],
    widths_cm=[1.5, 12.5, 1.5]
)
add_table_caption("Table 3.1: Functional Requirements (MoSCoW prioritisation).")

h3("3.3.2 Non-Functional Requirements")
add_table(
    ["ID", "Category", "Requirement"],
    [
        ["NFR-01", "Reproducibility", "The full pipeline shall be re-runnable end-to-end from raw data with fixed random seeds and versioned inputs."],
        ["NFR-02", "Performance", "A full pipeline run over the AOI and study period shall complete within practical limits on Colab/Kaggle free tiers."],
        ["NFR-03", "Data provenance", "Each dataset shall be documented with source, licence, resolution and checksum."],
        ["NFR-04", "Transparency", "Model decisions shall be explainable via feature importance or SHAP, and the HMM transition matrix shall be reported."],
        ["NFR-05", "Scalability", "The grid and feature pipeline shall support changing AOI extent and grid size without code changes."],
        ["NFR-06", "Portability", "The system shall run locally, in Google Colab and in Kaggle, and be containerisable via Docker."],
        ["NFR-07", "Maintainability", "Code shall be modular by pipeline layer and achieve at least 70% line coverage in automated tests."],
        ["NFR-08", "Integrity / ethics", "Outputs shall be framed as flood-pressure estimates, not confirmed flood predictions, with limitations stated."],
    ],
    widths_cm=[1.5, 3, 11]
)
add_table_caption("Table 3.2: Non-Functional Requirements.")

h2("3.4 System Architecture")
justified(
    "Figure 1 depicts the four-tier architecture: external data sources, a processing and "
    "feature-engineering tier, a modelling tier, and an output and evaluation tier. The external "
    "tier integrates CHIRPS and Meteo Rwanda rainfall, SRTM elevation, OpenStreetMap vectors and "
    "official flood-risk polygons from the Rwanda GeoPortal. The processing tier builds the "
    "spatial grid and engineers rainfall, terrain, hydrology and exposure features into a single "
    "grid-cell-by-day modelling table. The modelling tier hosts the supervised flood-pressure "
    "classifier and the HMM temporal layer, with trained artefacts kept in a model registry. The "
    "output tier produces risk maps, a notebook/dashboard, spatial validation against official "
    "polygons, and evaluation metrics including the estimated transition matrix."
)
add_figure(os.path.join(FIG, "fig01_architecture.png"),
           "Figure 1: System Architecture of the Geospatial Flood-Pressure Modelling System.",
           width_in=6.5)

h3("3.4.1 Data Definition and Acquisition")
justified(
    "The project uses real, openly available data sources, summarised in Table 3.3, prepared "
    "around the Kigali flood-risk corridor and runnable locally, in Google Colab or in Kaggle. "
    "The geolocation script queries OpenStreetMap through Nominatim and Overpass to locate the "
    "study-area boundary and download waterways, roads, buildings and public facilities, which "
    "are converted into distances, densities and exposure counts for each grid cell."
)
add_table(
    ["Dataset", "Purpose", "Role in the pipeline"],
    [
        ["Meteo Rwanda / ENACTS rainfall", "Rwanda-specific daily rainfall", "Primary source for daily and rolling rainfall totals"],
        ["CHIRPS daily rainfall", "Long, gridded daily rainfall (Climate Hazards Center, 2024)", "Strong fallback / cross-check for the rainfall series"],
        ["Official flood-risk polygons (Rwanda GeoPortal)", "Official spatial flood-risk information", "Label construction and spatial validation"],
        ["SRTM / NASADEM elevation (NASA JPL, 2020)", "Elevation for the study area", "Derive elevation and slope features"],
        ["OpenStreetMap data", "Rivers, roads, buildings, facilities", "Distance-to-river, road/building density, exposure"],
        ["Optional flood incident records (MINEMA)", "Strengthen validation if available", "Used only if records are credible, dated and geolocated"],
    ],
    widths_cm=[4.5, 5.5, 5.5]
)
add_table_caption("Table 3.3: Datasets Within the Scope of the Study.")
justified(
    "The final modelling table holds one row per grid cell per day, so that rainfall varies over "
    "time while terrain, river distance, road and building density and flood-polygon intersection "
    "remain attached to each spatial unit. Each dataset is documented with provenance, licence and "
    "any privacy constraints; raw files are stored with checksums to ensure reproducibility."
)

h3("3.4.2 Machine Learning and HMM Pipeline")
justified(
    "The pipeline is implemented in Python using geopandas, rasterio/rioxarray and shapely for "
    "geoprocessing, scikit-learn and XGBoost for classification, and hmmlearn for the temporal "
    "layer. It is organised in seven layers (Table 3.4) so that the role of each dataset and "
    "model is explicit and independently testable. The expected output is a per-cell-day "
    "predicted flood-pressure state with probability, plus an estimated daily transition matrix "
    "and next-step state probabilities."
)
add_table(
    ["Layer", "What happens", "Output"],
    [
        ["1. Study area & grid", "AOI converted into a grid of spatial cells", "grid_id, geometry, centroid coordinates"],
        ["2. Climate features", "Daily rainfall matched to cells; rolling totals computed", "rainfall_mm, rain_3d/7d/14d"],
        ["3. Geospatial features", "Elevation, slope, river distance, densities, polygon intersection", "Static spatial features per cell"],
        ["4. Label construction", "Flood-pressure states derived from rainfall + official risk", "low / moderate / high label"],
        ["5. ML classification", "Baseline and ML models classify pressure states", "Predicted state and probability"],
        ["6. HMM temporal layer", "State transitions estimated over the daily step", "Transition matrix; next-step probabilities"],
        ["7. Evaluation & mapping", "Compare with baselines and official polygons", "Metrics, maps and interpretation"],
    ],
    widths_cm=[3.2, 7.3, 5]
)
add_table_caption("Table 3.4: ML and HMM Pipeline Layers.")
justified(
    "The Hidden Markov Model is used strictly as the temporal layer. The hidden states are the "
    "latent flood-pressure conditions of a grid cell (low, moderate, high); the observed "
    "(emission) variables are rainfall totals, slope, elevation, distance to river, road and "
    "building density, and flood-polygon intersection. With a daily time step, the transition "
    "matrix represents one-day transitions, and multi-day tendencies can be approximated by "
    "matrix powers. Transition probabilities are estimated from data via the Baum-Welch algorithm "
    "and sequences decoded with Viterbi (Xie, Jiang, & Sainju, 2018; Stoner & Economou, 2020). If only weekly "
    "rainfall summaries were available, the model would be described as a weekly tendency model "
    "rather than a daily warning model, and claims reduced accordingly."
)

h3("3.4.3 Model Comparison and Architecture")
justified(
    "The framework is benchmarked against simple baselines so that the geospatial ML and HMM "
    "components must demonstrate value beyond rainfall thresholds and static risk maps. Defaults "
    "are explicitly avoided; final selection is justified by held-out classification performance, "
    "high-pressure recall and spatial agreement with official polygons."
)
add_table(
    ["Model / baseline", "Purpose", "Expected role / metric"],
    [
        ["Rainfall-threshold baseline", "Classify pressure from fixed rainfall totals only", "Tests whether geospatial features add value"],
        ["Static flood-polygon baseline", "Flag cells intersecting official risk polygons", "Tests whether temporal rainfall adds value"],
        ["Logistic Regression / Decision Tree", "Interpretable ML baseline", "Shows whether a basic learner classifies states"],
        ["Random Forest", "Robust ensemble on tabular geodata", "Macro F1 ~0.72-0.82; feature importance"],
        ["XGBoost", "Main gradient-boosted model", "Macro F1 ~0.75-0.85; SHAP explanations"],
        ["HMM temporal layer", "Model daily state transitions", "Next-step accuracy; sequence stability"],
    ],
    widths_cm=[4, 6, 5.5]
)
add_table_caption("Table 3.5: Candidate Models and Baselines for Flood-Pressure Modelling.")
justified(
    "Classification performance is reported with accuracy, macro F1-score, balanced accuracy, "
    "precision, recall and the confusion matrix, with particular weight on recall for the "
    "high-pressure state, since missing high-pressure areas is more serious than over-warning. "
    "Spatial agreement is measured as the overlap and the percentage of predicted high-pressure "
    "cells inside official flood-risk zones, and temporal performance via next-step accuracy and "
    "transition-matrix interpretation."
)

h2("3.5 UML and Modelling Diagrams")

h3("3.5.1 Use Case Diagram")
justified(
    "Figure 3 identifies four actors (Researcher/Analyst, Disaster Manager (MINEMA), Urban "
    "Planner and the Meteo Rwanda / data API) and the primary use cases, ranging from acquiring "
    "rainfall and geodata and building the spatial grid to running HMM transitions, generating "
    "risk maps and validating against official flood polygons."
)
add_figure(os.path.join(FIG, "fig03_use_case.png"),
           "Figure 3: Use Case Diagram.")

h3("3.5.2 Class Diagram")
justified(
    "Figure 4 specifies the principal domain classes and their relationships. A GridCell "
    "aggregates many RainfallObservation records and is linked to HydrologyExposure features and "
    "FloodPressureLabel records. A ClassifierModel produces predictions consumed by the HMMModel, "
    "which feeds the RiskMap and EvaluationReport, while ModelRun and Dataset capture provenance "
    "and reproducibility."
)
add_figure(os.path.join(FIG, "fig04_class.png"),
           "Figure 4: Class Diagram (UML).", width_in=6.5)

h3("3.5.3 Entity-Relationship Diagram")
justified(
    "Figure 5 shows the relational schema in Crow's-foot notation. The schema normalises grid "
    "cells, daily rainfall, hydrology and exposure features, flood-risk polygons, pressure "
    "states, model runs, transitions, datasets and evaluation metrics, with primary and foreign "
    "keys identified."
)
add_figure(os.path.join(FIG, "fig05_erd.png"),
           "Figure 5: Entity-Relationship Diagram (Crow's-Foot Notation).",
           width_in=6.5)

h3("3.5.4 Sequence Diagram")
justified(
    "Figure 6 traces the most important interaction: a daily flood-pressure estimation run, from "
    "the analyst's request through the pipeline orchestrator, data store, feature engine, "
    "classifier and HMM temporal layer, ending with a rendered risk map and evaluation report."
)
add_figure(os.path.join(FIG, "fig06_sequence.png"),
           "Figure 6: Sequence Diagram: Daily Flood-Pressure Estimation.",
           width_in=6.5)

h3("3.5.5 HMM State-Transition Model")
justified(
    "Figure 2 details the core contribution: the Hidden Markov Model over three hidden "
    "flood-pressure states (low, moderate, high). Directed edges represent daily transition "
    "probabilities and self-loops the persistence of a state, while the observed rainfall, "
    "terrain, hydrology and exposure variables form the emissions from which states are inferred. "
    "The transition and emission parameters are estimated from data rather than fixed."
)
add_figure(os.path.join(FIG, "fig02_hmm_states.png"),
           "Figure 2: Hidden Markov Model of Daily Flood-Pressure State Transitions.",
           width_in=6.5)

h2("3.6 Development Tools")
add_table(
    ["Layer", "Tool / Technology", "Justification"],
    [
        ["Language", "Python 3.11", "Dominant language for geospatial ML; rich ecosystem"],
        ["Geoprocessing", "geopandas, shapely, rasterio / rioxarray", "Vector and raster handling for grid and features"],
        ["Data access", "osmnx, Overpass / Nominatim, requests", "Reproducible OpenStreetMap and rainfall acquisition"],
        ["ML libraries", "scikit-learn, XGBoost, Optuna", "Baselines, gradient boosting and hyper-parameter tuning"],
        ["Temporal model", "hmmlearn", "Hidden Markov Model fitting (Baum-Welch) and decoding"],
        ["Explainability", "SHAP", "Transparent feature attribution (NFR-04)"],
        ["Mapping / viz", "matplotlib, folium / kepler.gl", "Static and interactive flood-pressure maps"],
        ["GIS desktop", "QGIS", "Inspection and cartographic finishing"],
        ["Notebooks / compute", "Jupyter, Google Colab, Kaggle", "Free-tier reproducible execution"],
        ["Dashboard", "Streamlit", "Inspect predictions by date and grid cell"],
        ["Version control / CI", "Git + GitHub, GitHub Actions", "Reproducibility and automated testing"],
        ["Containerisation", "Docker", "Portable deployment across providers"],
    ],
    widths_cm=[3.2, 5, 7.8]
)
add_table_caption("Table 3.6: Development Tools and Justification.")

# ---------------------------------------------------------------------------
# REFERENCES (with clickable hyperlinks)
# ---------------------------------------------------------------------------
h1("References")

# Each entry: (citation_text, url_or_None). The url is rendered as a clickable link.
# All references are limited to the 2018-2026 window.
refs = [
    ("Climate Hazards Center. (2024). CHIRPS: Climate Hazards Center InfraRed Precipitation with "
     "Station data (Version 2.0) [Data set]. University of California, Santa Barbara. ",
     "https://www.chc.ucsb.edu/data/chirps"),
    ("Feng, W., Yang, L. E., Ai, M., Chen, S., Wang, Z., Wu, W., Chen, J., Fang, Y., Xu, Y., & "
     "Garschagen, M. (2025). FRAMe: Empirically informed agent-based modeling of flood resilience "
     "in the Mekong River Basin. MethodsX, 15, 103682. ",
     "https://doi.org/10.1016/j.mex.2025.103682"),
    ("Hahirwabasenga, J., Nilsson, E., Larson, M., Bizimana, H., Wali, U. G., & Persson, M. "
     "(2024). Flooding in Sebeya catchment, Rwanda - A review of causes, impacts, and management. "
     "International Journal of Disaster Risk Reduction, 114, 105012. ",
     "https://doi.org/10.1016/j.ijdrr.2024.105012"),
    ("Intergovernmental Panel on Climate Change. (2022). Climate change 2022: Impacts, adaptation "
     "and vulnerability. Cambridge University Press. ",
     "https://www.ipcc.ch/report/ar6/wg2/"),
    ("Khosravi, K., Pham, B. T., Chapi, K., Shirzadi, A., Shahabi, H., Revhaug, I., ... Bui, D. T. "
     "(2018). A comparative assessment of decision trees algorithms for flash flood "
     "susceptibility modeling at Haraz watershed, northern Iran. Science of the Total "
     "Environment, 627, 744-755. ",
     "https://doi.org/10.1016/j.scitotenv.2018.01.266"),
    ("Long, L., Junjia, Y., Jiao, W., Alias, A. H., Haron, N. A., & Abu Bakar, N. (2024). Urban "
     "flood resilience evaluation in China: A systematic review of frameworks, methods, and "
     "limitations. Geomatics, Natural Hazards and Risk, 15(1), 2445631. ",
     "https://doi.org/10.1080/19475705.2024.2445631"),
    ("Ma, M., Zhao, G., He, B., Li, Q., Dong, H., Wang, S., & Wang, Z. (2021). XGBoost-based "
     "method for flash flood risk assessment. Journal of Hydrology, 598, 126382. ",
     "https://doi.org/10.1016/j.jhydrol.2021.126382"),
    ("Mind'je, R., Li, L., Amanambu, A. C., Nahayo, L., Nsengiyumva, J. B., Gasirabo, A., & "
     "Mindje, M. (2019). Flood susceptibility modeling and hazard perception in Rwanda. "
     "International Journal of Disaster Risk Reduction, 38, 101211. ",
     "https://doi.org/10.1016/j.ijdrr.2019.101211"),
    ("Mind'je, R., Li, L., Kayumba, P. M., Mindje, M., Ali, S., & Umugwaneza, A. (2021). "
     "Integrated geospatial analysis and hydrological modeling for peak flow and volume "
     "simulation in Rwanda. Water, 13(20), 2926. ",
     "https://doi.org/10.3390/w13202926"),
    ("Ministry in charge of Emergency Management [MINEMA]. (2023). Annual disaster effects report "
     "2023. Government of Rwanda. ",
     "https://www.minema.gov.rw/fileadmin/user_upload/Minema/Publications/Disaster_Effects/Annual_Disaster_Effects_Report_2023.pdf"),
    ("Mosavi, A., Ozturk, P., & Chau, K.-W. (2018). Flood prediction using machine learning "
     "models: Literature review. Water, 10(11), 1536. ",
     "https://doi.org/10.3390/w10111536"),
    ("NASA JPL. (2020). NASADEM merged DEM global 1 arc second V001 [Data set]. NASA Land "
     "Processes Distributed Active Archive Center. ",
     "https://doi.org/10.5067/MEASURES/NASADEM/NASADEM_HGT.001"),
    ("National Institute of Statistics of Rwanda [NISR]. (2023). Statistical yearbook 2023. "
     "Government of Rwanda. ",
     "https://www.statistics.gov.rw/"),
    ("Nzabonantuma, L., Nduwayezu, G., Naghibi, A., Nilsson, E., Wali, U. G., & Larson, M. "
     "(2025). Flood susceptibility mapping in the Nyabarongo Catchment, Rwanda, based on data "
     "analysis and modeling. Geomatics, Natural Hazards and Risk, 16(1), 2556987. ",
     "https://doi.org/10.1080/19475705.2025.2556987"),
    ("OpenStreetMap contributors. (2024). OpenStreetMap [Data set]. OpenStreetMap Foundation. ",
     "https://www.openstreetmap.org"),
    ("Papaioannou, G., Varlas, G., Terti, G., Papadopoulos, A., Loukas, A., Panagopoulos, Y., & "
     "Dimitriou, E. (2019). Flood inundation mapping at ungauged basins using coupled "
     "hydrometeorological-hydraulic modelling: The catastrophic case of the 2006 flash flood in "
     "Volos City, Greece. Water, 11(11), 2328. ",
     "https://doi.org/10.3390/w11112328"),
    ("Rwanda Meteorology Agency. (2022). Downscaled climate projections for Rwanda. Government of "
     "Rwanda. ",
     "https://www.meteorwanda.gov.rw/"),
    ("Stoner, O., & Economou, T. (2020). An advanced hidden Markov model for hourly rainfall time "
     "series. Computational Statistics & Data Analysis, 152, 107045. ",
     "https://doi.org/10.1016/j.csda.2020.107045"),
    ("The New Times. (2026, February 12). Kigali steps up drainage patrols after floods, crime "
     "incidents. The New Times Rwanda. ",
     "https://www.newtimes.co.rw/article/33240/news/security/kigali-steps-up-drainage-patrols-after-floods-crime-incidents"),
    ("vom Brocke, J., Hevner, A., & Maedche, A. (2020). Introduction to design science research. "
     "In J. vom Brocke, A. Hevner, & A. Maedche (Eds.), Design science research. Cases (pp. "
     "1-13). Springer. ",
     "https://doi.org/10.1007/978-3-030-46781-4_1"),
    ("Xie, M., Jiang, Z., & Sainju, A. M. (2018). Geographical hidden Markov tree for flood "
     "extent mapping. In Proceedings of the 24th ACM SIGKDD International Conference on Knowledge "
     "Discovery & Data Mining (pp. 2545-2554). Association for Computing Machinery. ",
     "https://doi.org/10.1145/3219819.3220053"),
]
for text, url in refs:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.left_indent = Cm(1.27)
    pf.first_line_indent = Cm(-1.27)
    r = p.add_run(text)
    set_run(r, size=11.5, font="Times New Roman")
    if url:
        add_hyperlink(p, url, url)

# ---------------------------------------------------------------------------
doc.save(OUT)
print("Saved:", OUT)
print("Size:", os.path.getsize(OUT), "bytes")
